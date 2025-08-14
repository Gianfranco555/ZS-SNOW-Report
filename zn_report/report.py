from __future__ import annotations

import base64
from pathlib import Path

import docx
from docx.shared import Inches
from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML

from .config import Config


def _create_docx_report(
    metrics: dict, chart_paths: dict[str, Path], config: Config
) -> docx.document.Document:
    """Create a DOCX report programmatically."""
    document = docx.Document()
    document.add_heading(config.title, level=1)

    # KPIs
    document.add_heading("Executive KPIs", level=2)
    kpis = metrics.get("kpis", {})
    for key, value in kpis.items():
        label = key if " " in key else key.replace("_", " ").title()
        p_text = f"{label}: {value}"
        if key == "avg_ttr_hours":
            p_text += " hours"
        document.add_paragraph(p_text)

    # Open by State
    open_by_state_data = metrics.get("tables", {}).get("open_by_state", [])
    if open_by_state_data:
        document.add_heading("Open by State", level=2)
        for item in open_by_state_data:
            document.add_paragraph(
                f"{item['state']} — {item['count']} ({item['percent']}%)",
                style="List Bullet",
            )

    # Charts
    document.add_heading("Charts", level=2)
    for chart_id, chart_path in chart_paths.items():
        document.add_heading(chart_id.replace("_", " ").title(), level=3)
        document.add_picture(str(chart_path), width=Inches(6.0))

    # Tables
    document.add_heading("Data Tables", level=2)
    tables = metrics.get("tables", {})
    for table_name, table_data in tables.items():
        if table_name == "open_by_state":
            continue

        document.add_heading(table_name.replace("_", " ").title(), level=3)
        if not table_data:
            document.add_paragraph("No data available.")
            continue

        headers = list(table_data[0].keys())
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"

        # Populate header
        for i, header in enumerate(headers):
            table.cell(0, i).text = header.replace("_", " ").title()

        # Populate rows
        for row_data in table_data:
            row_cells = table.add_row().cells
            for i, header in enumerate(headers):
                row_cells[i].text = str(row_data.get(header, ""))

        # Add total row if 'count' column exists
        if "count" in headers:
            total = sum(row.get("count", 0) for row in table_data)
            total_cells = table.add_row().cells
            count_idx = headers.index("count")

            # Handle case where 'count' is the first column
            if count_idx == 0:
                total_cells[0].text = f"Total: {total}"
                total_cells[0].paragraphs[0].runs[0].bold = True
            else:
                total_cells[0].text = "Total"
                total_cells[0].paragraphs[0].runs[0].bold = True
                total_cells[count_idx].text = str(total)
                total_cells[count_idx].paragraphs[0].runs[0].bold = True

    if config.branding.footer:
        section = document.sections[0]
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = config.branding.footer

    return document


def assemble_report(
    metrics: dict,
    chart_paths: dict[str, Path],
    config: Config,
    template_path: Path,
    output_path: Path,
):
    """Generate a report in PDF or DOCX format.

    The output format is determined by the file extension of `output_path`.
    """
    output_suffix = output_path.suffix.lower()

    if output_suffix == ".pdf":
        # Create data URIs to embed images directly into the HTML
        chart_data_uris = {}
        for key, path in chart_paths.items():
            with open(path, "rb") as image_file:
                encoded_string = base64.b64encode(image_file.read()).decode("utf-8")
            chart_data_uris[key] = f"data:image/png;base64,{encoded_string}"

        # Render HTML from Jinja2 template
        env = Environment(
            loader=FileSystemLoader(template_path.parent),
            enable_async=True,
        )
        template = env.get_template(template_path.name)
        rendered_html = template.render(
            metrics=metrics,
            charts=chart_data_uris,
            config=config,
        )

        # Generate PDF from HTML (no base_url needed)
        html = HTML(string=rendered_html)
        html.write_pdf(output_path)

    elif output_suffix == ".docx":
        document = _create_docx_report(metrics, chart_paths, config)
        document.save(output_path)

    else:
        raise ValueError(f"Unsupported output format: {output_suffix}")
