import os
import hashlib
from pathlib import Path

import docx
import pytest
import pandas as pd

from zn_report.config import Config, load_config
from zn_report.report import assemble_report
from zn_report.metrics import compute_metrics
from zn_report.charts import render_charts
from zn_report.config import Style, Palette

GOLDEN_DIR = Path(__file__).parent / "golden"

def get_file_hash(path: Path) -> str:
    """Computes the SHA256 hash of a file."""
    sha256 = hashlib.sha256()
    with open(path, "rb") as f:
        while chunk := f.read(8192):
            sha256.update(chunk)
    return sha256.hexdigest()

def check_golden_hash(generated_path: Path, test_name: str):
    """Checks a generated file's hash against its golden hash file."""
    GOLDEN_DIR.mkdir(exist_ok=True)
    golden_path = GOLDEN_DIR / f"{test_name}.sha256"
    generated_hash = get_file_hash(generated_path)

    if not golden_path.exists():
        golden_path.write_text(generated_hash)
        pytest.skip(f"Golden file '{golden_path.name}' created. Please review and commit.")

    expected_hash = golden_path.read_text().strip()
    assert generated_hash == expected_hash, f"Hash mismatch for {test_name}"

@pytest.fixture(scope="session")
def sample_config() -> Config:
    """Return a sample Config object."""
    return load_config({"title": "Test Report", "branding": {"footer": "Test Footer"}})

@pytest.fixture(scope="session")
def processed_metrics(shared_df: pd.DataFrame) -> dict:
    """Provides a sample metrics dictionary derived from the shared DataFrame."""
    return compute_metrics(shared_df, start="2025-07-01", end="2025-07-31", tz="UTC")

@pytest.fixture(scope="session")
def chart_style() -> Style:
    """Provides a default Style object for tests."""
    return Style(palette=Palette(primary="#000000", secondary="#444444", accent="#ff0000", muted="#888888", categorical=["#ff0000", "#00ff00", "#0000ff"]), font_family="sans-serif")

@pytest.fixture
def sample_chart_paths(processed_metrics: dict, chart_style: Style, tmp_path: Path) -> dict[str, Path]:
    """Create real chart files based on processed metrics."""
    return render_charts(processed_metrics, chart_style, tmp_path)

def _verify_docx(path: Path, metrics: dict):
    """Verify the content of the generated DOCX file."""
    doc = docx.Document(path)
    # Combine all paragraph and table text for easier searching
    all_text = "\n".join([p.text for p in doc.paragraphs])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += "\n" + cell.text

    # Verify title
    assert doc.paragraphs[0].text == "Test Report"
    assert doc.paragraphs[0].style.name.startswith("Heading")

    # Verify KPIs by checking for the label and the value in the document text.
    # This is more robust to formatting changes in the template.
    kpis = metrics["kpis"]
    assert "Resolved Count" in all_text
    assert str(kpis['resolved_count']) in all_text

    assert "Avg Ttr Hours" in all_text
    assert str(kpis['avg_ttr_hours']) in all_text

    # Verify presence of table titles (which are also chart titles)
    assert "Resolved By Assignee" in all_text
    assert "Top 5 Tags" in all_text
    assert "Open By State" in all_text

    # Verify image count matches charts generated
    # There are 5 charts defined in charts.py
    image_count = len(doc.inline_shapes)
    assert image_count == 5

def test_assemble_report(
    tmp_path: Path,
    sample_config: Config,
    processed_metrics: dict,
    sample_chart_paths: dict[str, Path],
):
    """Test the report assembly for both PDF and DOCX."""
    # Set SOURCE_DATE_EPOCH for deterministic PDF generation
    os.environ["SOURCE_DATE_EPOCH"] = "1672531200"  # 2023-01-01
    template_path = Path("zn_report/templates/report.html.j2")

    # --- Test PDF Generation ---
    pdf_output_path = tmp_path / "report.pdf"
    assemble_report(
        metrics=processed_metrics,
        chart_paths=sample_chart_paths,
        config=sample_config,
        template_path=template_path,
        output_path=pdf_output_path,
    )
    assert pdf_output_path.exists()
    check_golden_hash(pdf_output_path, "report.pdf")

    # --- Test DOCX Generation ---
    docx_output_path = tmp_path / "report.docx"
    assemble_report(
        metrics=processed_metrics,
        chart_paths=sample_chart_paths,
        config=sample_config,
        template_path=template_path,
        output_path=docx_output_path,
    )
    assert docx_output_path.exists()
    _verify_docx(docx_output_path, processed_metrics)

    # Clean up environment variable
    del os.environ["SOURCE_DATE_EPOCH"]
